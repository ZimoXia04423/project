# -*- coding: utf-8 -*-
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document


def main():
    path = Path(__file__).resolve().parents[1] / "短视频.docx"
    doc = Document(str(path))
    for p in doc.paragraphs:
        if "系统还实现了动态词典更新机制" in p.text:
            p.clear()
            p.add_run(
                "TF-IDF 算法实现基于 jieba.analyse.extract_tags 接口，其 IDF 统计来自 jieba 内置语料，"
                "对合并后的标题或评论长文本抽取得分最高的 Top-K 词。相比词频，TF-IDF 能抑制在全部文档中普遍出现的高频词、突出区分性词；"
                "本实现未包含「定期热词更新自定义词典」的批处理，若需领域适配可在未来将 jieba.load_userdict 等步骤纳入采集流水线。"
            )
            break
    doc.save(str(path))
    print("OK")


if __name__ == "__main__":
    main()
