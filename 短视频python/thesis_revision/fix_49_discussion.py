# -*- coding: utf-8 -*-
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document


def main():
    path = Path(__file__).resolve().parents[1] / "短视频.docx"
    doc = Document(str(path))

    for para in doc.paragraphs:
        if para.text.startswith("综合以上功能测试及 4.6～4.8 节的方法学补充"):
            para.clear()
            para.add_run(
                "综合以上功能测试及 4.6～4.8 节的方法学补充，可以得出以下几点发现："
                "第一，B 站热门排行榜中的视频呈现分区集中效应，生活、游戏和娱乐类内容占比较高，反映用户内容消费偏好。"
                "第二，热度演化多符合「快速增长—逐渐饱和」的传播特征，但峰值与增速因 UP 主影响力、题材与发布时间差异较大。"
                "第三，关键词分析揭示热点主题的迁移规律；词频与 TF-IDF、标题源与评论源的对照结论与 4.7 节一致，有助于区分「议题命名」与「互动用语」。"
                "第四，情感分布整体偏正面（见 4.5），但 4.6 节提醒应对 SnowNLP 误判与弹幕语境保持审慎。"
            )
            break

    doc.save(str(path))
    print("OK")


if __name__ == "__main__":
    main()
