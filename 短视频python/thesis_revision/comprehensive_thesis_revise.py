# -*- coding: utf-8 -*-
"""
《短视频.docx》全面修订：统一第3章目录与正文、口径与实现；加深3.8可视化链路；夯实第4章。
"""
import sys
import shutil
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def replace_para_if_contains(old_substring, new_text, doc):
    for para in doc.paragraphs:
        if old_substring in para.text:
            para.clear()
            para.add_run(new_text)
            return True
    return False


def set_heading_text(paragraph, new_title, level_style="Heading 2"):
    paragraph.clear()
    paragraph.add_run(new_title)
    try:
        paragraph.style = level_style
    except Exception:
        pass


def main():
    root = Path(__file__).resolve().parents[1]
    path = root / "短视频.docx"
    draft = root / "短视频_修订稿.docx"
    # 若修订稿已存在（上轮因占用另存），在其基础上迭代，避免重复插入导读段落
    src = draft if draft.exists() else path
    backup = root / "短视频_修订前备份.docx"
    if path.exists() and not backup.exists():
        shutil.copy2(path, backup)

    doc = Document(str(src))

    # --- 3.1 下插入 3.1.1 标题（若不存在）---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "3.1 系统总体架构设计":
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if nxt is not None and nxt.text.strip().startswith("3.1.1"):
                break
            insert_paragraph_after(
                p,
                "3.1.1 数据口径与模块映射（设计与实现一致性说明）",
                "Heading 3",
            )
            break

    # 章节目录纠偏：3.2 实为数据库设计
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "3.2 数据抓取模块设计与实现" or t.startswith("3.2 数据抓取模块设计与实"):
            set_heading_text(para, "3.2 数据库设计", "Heading 2")
            break

    # 在 3.1 总体分层段之后插入本章阅读路径（避免重复插入）
    roadmap_marker = "本章各节阅读路径"
    if not any(roadmap_marker in p.text for p in doc.paragraphs):
        for i, p in enumerate(doc.paragraphs):
            if "分层模块化架构" in p.text:
                insert_paragraph_after(
                    p,
                    "【本章阅读路径】" + roadmap_marker + "：3.2 给出库表与索引（存储前提）→3.3 说明采集/处理/分析的职责划分（设计视角）→3.4 说明 Web 与内部调用关系→3.5 对应爬虫与 HTTP 实现→3.6 对应 cleaner 清洗与类型规整→3.7 对应 analyzer 算法→3.8 对应可视化与异步图表→3.9 汇总文件映射。"
                    "其中「数据清洗」在 3.3.2 表述设计原则，在 3.6 与 3.9 给出与 cleaner/data_cleaner.py 一致的步骤，避免重复时以 3.1.1 口径为准。",
                )
                break

    # --- 正文事实对齐 ---
    for para in doc.paragraphs:
        tx = para.text
        if "DataCleaner类" in tx and "clean_videos" in tx:
            para.clear()
            para.add_run(
                "数据处理模块对应 cleaner/data_cleaner.py：以函数 clean_videos、clean_comments、prepare_ranking_records 组织，"
                "基于 Pandas 完成 HTML 去标签、去重、缺失填充与数值转换；本文未实现独立的 DataCleaner 类，亦无单独的 validate_data()——"
                "质量控制主要体现在清洗规则（非空、类型可解析）与入库前字段一致性上。"
            )
        if "各个分析服务之间保持松耦合关系" in tx and "HotRankingService" in tx:
            para.clear()
            para.add_run(
                "分析逻辑以函数为单位集中在 analyzer/data_analyzer.py（analyze_hot_ranking、analyze_trend、analyze_keywords、analyze_keywords_tfidf、analyze_sentiment），"
                "由 Flask 路由直接调用；本文未拆分为 HotRankingService、TrendAnalysisService 等独立服务类，亦未在模块级实现统一的性能监控或结果缓存接口。"
            )
        if "分析服务接口提供" in tx and "/api/analysis/" in tx:
            para.clear()
            para.add_run(
                "业务页面以 GET 路由为主（如 /ranking、/trend、/keywords、/sentiment），在视图函数内直接调用分析函数并渲染模板；"
                "图表异步加载走 /api/chart/ranking、/api/chart/keywords、/api/chart/sentiment、/api/chart/trend，返回 Pyecharts 嵌入片段或 Matplotlib 静态图所需数据。"
                "本文未另行实现 /api/analysis/* 风格的独立分析 REST 资源层。"
            )
        if "模块间的数据传输采用标准化的数据传输对象（DTO）模式" in tx:
            para.clear()
            para.add_run(
                "层间数据以 Python dict / list 及 PyMySQL 游标结果传递，数据库字段名与爬虫字典键保持一致；"
                "未引入 VideoDTO、CommentDTO 等显式 DTO 类型，亦未定义基于 BaseException 的分层异常体系或 Mock 注入点。"
            )
        if "频率控制实现采用令牌桶算法" in tx:
            para.clear()
            para.add_run(
                "频率控制由 SCRAPER_CONFIG[\"request_interval\"] 在评论分页之间以及「视频与视频」抓取之间 sleep 实现；"
                "HTTP 层失败由 _request_with_retry 按固定 retry_delay 重试至多 max_retries 次，而非令牌桶算法。"
            )
        if "反爬虫机制的实现包含多个维度的伪装策略" in tx and "代理池" in tx:
            para.clear()
            para.add_run(
                "请求侧主要使用 config.HEADERS 中的 User-Agent、Referer 等静态头；run_full_crawl 通过 request_interval 限速。"
                "代码中未实现动态 User-Agent 列表、代理池轮换与会话伪装登录；合规采集以低频次与重试为主。"
            )
        if "为了提高采集效率和可靠性，模块还实现了智能重试机制和断点续传功能" in tx:
            para.clear()
            para.add_run(
                "为提高可靠性，HTTP 请求通过 _request_with_retry 在失败时按固定间隔重试至多 max_retries 次。"
                "run_full_crawl 未实现断点续传或任务 checkpoint；大批量采集需在一次运行内完成或由外部调度重复执行。"
            )
        if "指数退避重试策略" in tx or ("重试间隔逐渐增长" in tx and "3秒、6秒、12秒" in tx):
            para.clear()
            para.add_run(
                "重试等待时间为配置中的固定 retry_delay（秒），各次重试间隔相同，直至用尽 max_retries。"
            )
        if "数据验证实现采用多级验证机制" in tx and "批次级验证" in tx:
            para.clear()
            para.add_run(
                "入库前的有效性主要体现在清洗阶段：非空、可解析数值、键完整；代码未单独实现批次级统计血缘与自动化质量报表。"
                "若需工业级质控，可在库内增加校验表或离线任务对 view_count 等做分位数告警。"
            )

    # 3.6 节首补充：设计—实现对应关系
    sect_36 = "3.6 数据处理模块实现"
    bridge_36 = "【与 3.3.2、3.9 的统一说明】"
    if not any(bridge_36 in p.text for p in doc.paragraphs):
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == sect_36:
                insert_paragraph_after(
                    p,
                    bridge_36
                    + "数据清洗「设计」体现在 3.3.2 的职责描述；「实现」唯一落在 cleaner/data_cleaner.py。"
                    "3.6.1～3.6.3 按清洗→类型转换与标准化→基本合法性约束的顺序展开，并与 3.1.1 去重键一致；"
                    "入库写入见 database/db_operations.py，表结构见 3.2。",
                )
                break

    # 3.8.1 增强：图表生成与界面交互链（追加深度，避免重复则前置检测）
    chart_deep = "【图表生成链路与界面深度说明】"
    if not any(chart_deep in p.text for p in doc.paragraphs):
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "3.8.1 图表生成引擎实现":
                # 找该小节第一个正文段落后追加
                j = i + 1
                while j < len(doc.paragraphs) and doc.paragraphs[j].style.name.startswith("Heading"):
                    j += 1
                if j < len(doc.paragraphs):
                    anchor = doc.paragraphs[j]
                    insert_paragraph_after(
                        anchor,
                        chart_deep
                        + "① 首屏：路由函数（app.py）调用 analyze_* 得到结构化数据，Jinja2 模板将轴标签、频次等序列化为 JSON，"
                        "由浏览器 ECharts 初始化本地柱状图/饼图（减少首屏阻塞）。② 异步增强：用户点击「加载图表」等控件后，前端通过 fetch 请求 /api/chart/*，"
                        "服务端再次调用 analyze_* 与 chart_generator：Pyecharts 返回 render_embed 的 HTML 片段，Matplotlib 返回 Base64 PNG Data URI，"
                        "前端注入 DOM 或赋值 img.src，实现「同指标、双引擎」对照。③ 趋势页：/api/chart/trend 仅在选定 bvid 后加载，避免空查询。"
                        "④ 语义约束：趋势曲线若含「★」「预测」等时间轴标记，应与 3.7.2/3.1.1 所述「补全/外推」一致解读。",
                    )
                break

    # 第4章 4.6 增补：显式「样本—标注—比对」句式（若尚无表格提示）
    tbl_hint = "【结果呈现建议】"
    if not any(tbl_hint in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if para.text.strip() == "4.6 情感模型有效性评估与标注实验":
                insert_paragraph_after(
                    para,
                    tbl_hint
                    + "建议在附录给出「标注细则 + 混淆矩阵（表4-A）+ 宏平均 Precision/Recall/F1」三件套："
                    "混淆矩阵行列为人工标签×模型预测；Accuracy 为对角线之和除以 N；"
                    "Positive 类的 Precision = TP/(TP+FP)。正文可节录矩阵与总体指标，完整表格放附录以保持版面整洁。",
                )
                break

    out_primary = path
    out_alt = root / "短视频_修订稿.docx"
    try:
        doc.save(str(out_primary))
        print("OK 修订已写入:", out_primary)
    except PermissionError:
        doc.save(str(out_alt))
        print("WARN: 原文件被占用，已写入:", out_alt, "（请关闭 Word 后将本文件重命名为 短视频.docx 覆盖即可）")
    if backup.exists():
        print("备份:", backup)


if __name__ == "__main__":
    main()
