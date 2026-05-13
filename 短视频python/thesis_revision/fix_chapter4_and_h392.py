# -*- coding: utf-8 -*-
"""修正 3.9.2 标题；删除错序的第4章插入块并按 4.6→4.7→4.8→4.9 重插。"""
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def delete_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def main():
    root = Path(__file__).resolve().parents[1]
    path = root / "短视频.docx"
    doc = Document(str(path))

    # 修正 3.9.2 标题
    for p in doc.paragraphs:
        if p.text.strip() == "3.5 数据可视化模块设计与实现":
            p.clear()
            p.add_run("3.9.2 可视化与异步接口在代码中的落点")
            break

    # 待删除的段落文本特征（错序插入批次）
    delete_markers = [
        "复杂查询可在 SQL 层完成多表联结",
        "短板方面：每次请求新建 PyMySQL",
        "示例：按 crawl_time 窗口 JOIN",
        "数据库采用 bilibili_hotspot（utf8mb4）",
        "4.8 数据库系统架构、短板与复杂分析查询",
        "实验观察到：TF-IDF 更抑制通用高频词",
        "在固定时间窗与数据源（标题/评论）下",
        "词频法：分词后计数，反映语言表层共现",
        "4.7 关键词提取方法对比与进一步验证",
        "该评估用于论证 SnowNLP 在本数据集上的可用边界",
        "在标注集上可计算总体准确率 Accuracy",
        "结果比对：以人工标签为 gold standard",
        "抽样与标注：从 comments 表随机分层抽取",
        "情感模型在本系统中指 SnowNLP 输出的连续得分",
        "4.6 情感模型有效性评估与标注实验",
    ]

    # 多次遍历删除（删除后索引变化）
    changed = True
    while changed:
        changed = False
        for p in list(doc.paragraphs):
            t = p.text.strip()
            for m in delete_markers:
                if t.startswith(m) or t == m:
                    delete_paragraph(p)
                    changed = True
                    break

    # 找到 4.5 最后一节正文（图4-5 之后的那段结论）
    anchor = None
    for p in doc.paragraphs:
        if "知识类和生活类视频的正面评论比例通常高于争议性话题的视频" in p.text:
            anchor = p
            break
    if anchor is None:
        print("ERROR: anchor not found")
        return

    # 自顶向下插入（每次插在 anchor 后，下一段插在上一段后）
    cur = anchor
    blocks = [
        ("Heading 2", "4.6 情感模型有效性评估与标注实验"),
        (
            "Normal",
            "情感模型在本系统中指 SnowNLP 输出的连续得分经阈值离散化后的三分类结果。"
            "记准类样本上的准确率（Accuracy）为正确条数占比；对单一类别（如正面）可计算精准率（Precision）"
            "= 该类别预测且确实为该类的条数 / 预测为该类的总条数。",
        ),
        (
            "Normal",
            "抽样与标注：从 comments 表随机分层抽取不少于 200 条（可按粗筛的正负中性占比分层），"
            "由标注员依据统一细则标注「正面、中性、负面」；存在分歧时由第三人裁定，以保证可复现。",
        ),
        (
            "Normal",
            "结果比对：以人工标签为 gold standard，与 SnowNLP 阈值分类逐条对照，统计混淆矩阵；"
            "可讨论中性类与相邻类混淆、反讽句误判等现象，并与 4.5 节的总体占比相互印证。",
        ),
        (
            "Normal",
            "在标注集上计算 Accuracy；对关注的类别（如正面）补充 Precision；多分类可补充宏平均 F1，"
            "用于说明模型在本数据集上的可用边界，而非重新训练 SnowNLP。",
        ),
        (
            "Normal",
            "该评估用于论证 SnowNLP 在本系统中作为探索性情感画像工具的可行性：即便整体分布合理，"
            "仍应对网络用语、讽刺与混合情感句保持审慎解释。",
        ),
        ("Heading 2", "4.7 关键词提取方法对比与进一步验证"),
        (
            "Normal",
            "词频法：分词后计数，反映表层共现；TF-IDF：Term Frequency × Inverse Document Frequency，"
            "压低全体语料高频通用词、抬高区分度词（实现为 jieba.analyse.extract_tags）。",
        ),
        (
            "Normal",
            "在相同时间窗与数据源（标题或评论）下，并行调用 analyze_keywords 与 analyze_keywords_tfidf，"
            "对比 Top-30 重叠率与互补词项，用以验证第 3 章两条路线在实际语料上的差异，而非重新训练分词模型。",
        ),
        (
            "Normal",
            "实验性结论：TF-IDF 更抑制「视频、哈哈」等通用高频词，有利于突出话题特征词；词频更直观但易受口语赘词支配。"
            "标题与评论分别提取时，前者偏议题命名，后者偏互动用语，二者 Top-K 重叠率可作为话题一致性的旁证。",
        ),
        ("Heading 2", "4.8 数据库系统架构、短板与复杂分析查询"),
        (
            "Normal",
            "数据库采用 bilibili_hotspot（utf8mb4），核心表 videos、ranking_records、comments，由 db_operations 读写；"
            "查询入口含 get_top_videos、query_videos、query_all_comments、get_video_trend 等，对应排行榜、关键词语料与趋势序列。",
        ),
        (
            "Normal",
            "复杂 SQL 示例：按 crawl_time 窗口联结 ranking_records 与 videos，还原榜单快照时刻的播放与分区分布；"
            "按 tname 分组求 AVG(view_count)；对 comments 按 bvid 聚合 COUNT 再与 videos 联结分析互动强度。",
        ),
        (
            "Normal",
            "短板：请求级新建 PyMySQL 连接，并发升高时开销明显；跨分区、多关键词布尔检索更多依赖应用层 pandas/jieba，"
            "库内未建全文索引；情感标签不落库，重复分析会重复调用 SnowNLP。",
        ),
        (
            "Normal",
            "改进方向：情感结果物化表或离线批处理；连接池；必要时引入 Elasticsearch 等支撑检索型分析。",
        ),
    ]

    for style, text in blocks:
        cur = insert_after(cur, text, style if style != "Normal" else None)

    doc.save(str(path))
    print("OK")


if __name__ == "__main__":
    main()
