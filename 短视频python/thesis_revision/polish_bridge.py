# -*- coding: utf-8 -*-
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document


def main():
    path = Path(__file__).resolve().parents[1] / "短视频.docx"
    doc = Document(str(path))

    for para in doc.paragraphs:
        if para.text.startswith('数据可视化模块采用"前端ECharts直接渲染'):
            para.clear()
            para.add_run(
                "承接 3.8 对可视化组件的分节说明，本节从「路由—查询—渲染」闭环归纳与代码文件的对应关系："
                "数据可视化模块采用「前端 ECharts 直接渲染 + 后端 Pyecharts 服务端片段 + Matplotlib 静态图」三层策略，"
            )
            para.add_run(
                "用户请求经 app.py 进入对应页面或 /api/chart/*，数据来自 data_analyzer 与数据库查询，"
                "图表由 chart_generator 输出后再注入模板。"
            )
            break

    for para in doc.paragraphs:
        if para.text.startswith("本系统的开发与测试环境配置如下"):
            para.clear()
            para.add_run(
                "本系统的开发与测试环境配置如下：操作系统为 Windows 10/11，Python 版本为 3.11，数据库为 MySQL 8.0，Web 浏览器为 Chrome 120。"
                "主要依赖库及版本包括 Flask、Requests、Pandas、jieba、SnowNLP、Pyecharts、Matplotlib 和 PyMySQL 等，所有依赖均通过 requirements.txt 统一管理。"
                "系统运行于本地开发服务器。除各功能页面的现象级测试外，本章在 4.6～4.8 节从情感标注评估、关键词方法对照与数据库分析能力三个方面补充方法学讨论。"
            )
            break

    for para in doc.paragraphs:
        if para.text.startswith("综合以上四个维度的分析结果"):
            para.clear()
            para.add_run(
                "综合以上功能测试及 4.6～4.8 节的方法学补充，可以得出以下几点发现：第一，B站热门排行榜中的视频呈现出明显的分区集中效应，"
                "生活、游戏和娱乐类内容占据了排行榜的主要位置，反映了平台用户的内容消费偏好。第二，热门视频的热度变化普遍符合「快速增长—逐渐饱和」的传播曲线特征，"
                "但不同视频的增长速度和峰值差异较大，受内容质量、发布时间和 UP 主影响力等多重因素影响。第三，关键词分析揭示了不同时间段热点话题的主题迁移规律，通过"
            )
            break

    doc.save(str(path))
    print("OK")


if __name__ == "__main__":
    main()
