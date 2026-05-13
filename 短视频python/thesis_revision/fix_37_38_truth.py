# -*- coding: utf-8 -*-
"""修正 3.7.1、3.7.2、3.8.1、3.7.4 中与代码不符的段落。"""
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document


def replace_first_contains(doc, needle, new_text):
    for para in doc.paragraphs:
        if needle in para.text:
            para.clear()
            para.add_run(new_text)
            return True
    return False


def main():
    path = Path(__file__).resolve().parents[1] / "短视频.docx"
    doc = Document(str(path))

    replace_first_contains(
        doc,
        "热度排名算法实现支持多指标综合排序和单指标专项排序两种模式",
        "热度排名与 get_top_videos、analyze_hot_ranking 一致：在 MySQL 层面对白名单内单一指标列（view_count、like_count、coin_count、danmaku_count、favorite_count、reply_count 之一）执行 ORDER BY 降序，再 LIMIT 取前 N 条，并可按 crawl_time 做时间窗过滤。"
        "本文未实现多指标加权求和为「综合热度分」的算法；若需扩展，可在应用层对多列线性加权后再排序，与当前实现相区别。",
    )

    replace_first_contains(
        doc,
        "单指标专项排序直接按照用户选择的指标进行降序排列，支持播放量",
        "指标选择由 Web 参数传入，服务端校验列名后拼入 SQL；LIMIT 与界面「显示条数」一致。并列时顺序依数据库实现而定，未强制二级排序键。",
    )

    replace_first_contains(
        doc,
        "当快照条数不超过 2 时，analyze_trend 启用基于 S 型曲线的历史补点与未来外推（实现见 data_analyzer.py）。",
        "当 get_video_trend 返回不少于 3 个时间点时，直接以实测序列绑定折线图。当返回不超过 2 个时间点时，为便于展示，analyze_trend 以 S 型增长曲线 _simulate_growth_curve 补全历史区段，并以 _predict_future 在末次观测值附近做简单外推作为「预测」示意；图例与轴标签上应区分补全点与真实点。",
    )

    replace_first_contains(
        doc,
        "预测趋势分析采用简单移动平均和指数平滑两种算法",
        "上述外推并非经典时间序列中的移动平均或指数平滑，也未在代码中计算相关系数、决定系数或置信区间；其定位是数据稀疏时的可视化辅助。",
    )

    replace_first_contains(
        doc,
        "图表生成引擎采用多引擎并行架构，集成了ECharts、Pyecharts和Matplotlib三种图表生成技术，为不同的应用场景提供最优的可视化解决方案",
        "图表呈现采用三种互补方式：模板中注入 JSON 由浏览器 ECharts 渲染；Pyecharts 在服务端生成 embed 片段经 /api/chart/* 返回；Matplotlib 输出 PNG 的 Data URI。对应 app.py 与 chart_generator.py 中的具体函数，满足交互与版式不同需求。",
    )

    replace_first_contains(
        doc,
        "图表生成引擎实现了统一的配置管理和样式系统，通过ChartConfig类定义标准的图表配置参数",
        "颜色、字号、图例等在各生成函数内联设置；未使用独立的 ChartConfig 类，亦未实现引擎级结果缓存。",
    )

    replace_first_contains(
        doc,
        "算法实现还包含文本预处理优化，针对B站评论的特点进行了专门改进。表情符号处理方面",
        "实现上逐条对评论文本调用 SnowNLP 取 sentiments，按 3.1.1 节阈值打标签；异常文本跳过。停用词等处理体现在关键词模块的 jieba 路径，情感函数内未再维护表情/网络语规则字典。",
    )

    doc.save(str(path))
    print("OK")


if __name__ == "__main__":
    main()
