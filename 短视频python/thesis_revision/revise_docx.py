# -*- coding: utf-8 -*-
"""
修订《短视频.docx》：统一第3章口径与编号，纠正与代码不一致处，充实第4章。
"""
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def replace_first(pred, new_text, doc):
    for para in doc.paragraphs:
        if para.text.strip() and pred(para.text):
            para.clear()
            para.add_run(new_text)
            return True
    return False


def main():
    root = Path(__file__).resolve().parents[1]
    path = root / "短视频.docx"
    doc = Document(str(path))

    # ---- 3.1 后插入 3.1.1（接在「3.1 系统总体架构设计」标题段之后）----
    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "3.1 系统总体架构设计":
            anchor_idx = i
            break
    if anchor_idx is None:
        print("ERROR: 找不到 3.1 标题")
        return

    anchor = doc.paragraphs[anchor_idx]
    h = insert_paragraph_after(
        anchor,
        "3.1.1 数据口径与模块映射（设计与实现一致性说明）",
        "Heading 3",
    )
    p2 = insert_paragraph_after(
        h,
        "为避免「总体设计」与「程序实现」各行其是，本文约定以下口径，全文分析、测试与界面展示均以此为准。"
        "（1）时间戳 crawl_time：一次排行榜抓取任务统一的采集时刻，视频表与评论表均携带该字段。"
        "（2）去重键：视频清洗以 (bvid, crawl_time) 去重；评论以 rpid 去重。"
        "（3）字符集：MySQL 采用 utf8mb4。"
        "（4）热度排行：由数据库查询按用户选择的单一指标字段 ORDER BY DESC 取 Top-N，本文程序未实现多指标加权合成。"
        "（5）趋势：当同一视频快照不超过 2 条时，analyze_trend 使用 S 型曲线补历史并用简单增长率外推预测，可视化时需区分实测与补全。"
        "（6）情感：SnowNLP 得分区间 [0,1]，阈值：>0.6 正面、<0.4 负面、其余中性。"
        "（7）关键词：词频路径为 jieba 分词 + 停用词 + Counter；TF-IDF 路径为 jieba.analyse.extract_tags。",
    )
    insert_paragraph_after(
        p2,
        "主数据流：scraper → data_cleaner → db_operations(MySQL) → data_analyzer → Flask(app.py) 路由与模板 → Pyecharts/Matplotlib/前端 ECharts。"
        "下列各节（3.2~3.8）按链路展开；若与本节口径冲突，以本节为准。",
    )

    # ---- 正文替换 ----
    reps = [
        (
            lambda t: "模块采用面向对象的设计方法，封装了BilibiliScraper类" in t,
            "数据采集实现位于 scraper/bilibili_scraper.py，以 fetch_ranking、fetch_video_comments、run_full_crawl 等函数组织，"
            "配合 config 中的接口地址与 SCRAPER_CONFIG（超时、重试间隔、每视频评论页数）。",
        ),
        (
            lambda t: "AnalysisService类" in t,
            "分析能力集中于 analyzer/data_analyzer.py 中的函数 analyze_hot_ranking、analyze_trend、analyze_keywords、analyze_keywords_tfidf、analyze_sentiment，由 Flask 直接调用。",
        ),
        (
            lambda t: "GET /api/videos用于获取视频列表" in t,
            "页面路由（如 /ranking、/keywords）直接调用分析函数；异步图表由 /api/chart/ranking、/api/chart/keywords、/api/chart/sentiment、/api/chart/trend 提供 Pyecharts HTML 或静态图。",
        ),
        (
            lambda t: "IDataRepository接口" in t,
            "内部通过 database/db_operations 中的函数访问 MySQL，未引入 IDataRepository 抽象接口层。",
        ),
        (
            lambda t: "频率控制实现采用令牌桶算法",
            "频率控制采用配置的 request_interval：评论分页请求之间及视频之间 sleep；HTTP 失败则由 _request_with_retry 按 retry_delay 重试至多 max_retries 次。",
        ),
        (
            lambda t: "指数退避重试策略，首次重试间隔为3秒",
            "重试间隔为配置中的固定 retry_delay（秒），直至用尽 max_retries。",
        ),
        (
            lambda t: "多指标综合排序采用加权求和算法",
            "热度排序实现为：get_top_videos 在允许字段集合内对单一列 ORDER BY DESC；未在库内计算多指标加权和。",
        ),
        (
            lambda t: "当数据点充足（≥5个时间点）时",
            "当某 bvid 的采集快照不少于 3 个时",
        ),
        (
            lambda t: "当数据点不足时，算法启用数据增强机制" in t and "趋势分析" in t,
            "当快照条数不超过 2 时，analyze_trend 启用基于 S 型曲线的历史补点与未来外推（实现见 data_analyzer.py）。",
        ),
        (
            lambda t: "预测趋势分析采用简单移动平均和指数平滑两种算法",
            "预测段采用 _predict_future 对末次观测值做简单递增外推，用于趋势页展示，并非经典移动平均或指数平滑模型。",
        ),
        (
            lambda t: "集成了ECharts、Pyecharts和Matplotlib三种图表生成技术，为不同的应用场景提供最优",
            "前端 ECharts 绑定模板内 JSON；服务端 Pyecharts 生成可嵌入 HTML；Matplotlib 输出 Base64 PNG，三者在 chart_generator.py 与页面中分工协作。",
        ),
        (
            lambda t: "通过ChartConfig类定义标准的图表配置参数",
            "图表样式在 chart_generator.py 各生成函数中内联配置，未单独定义 ChartConfig 类。",
        ),
    ]
    for pred, nt in reps:
        replace_first(pred, nt, doc)

    replace_first(
        lambda t: t.startswith("数据抓取模块是整个系统的数据源头"),
        "【与 3.2、3.5 呼应】下列段落给出文件级映射，避免重复编排小节号；技术口径见 3.1.1。",
        doc,
    )

    # 重复小节标题改名
    for para in doc.paragraphs:
        s = para.text.strip()
        if s.startswith("3.3数据清洗与存储模块设计与实现"):
            para.clear()
            para.add_run("3.9 文件级实现映射与数据流小结")
            para.style = "Heading 2"
        elif s == "3.4数据分析模块设计与实现":
            para.clear()
            para.add_run("3.9.1 分析模块在代码中的落点")
            para.style = "Heading 3"
        elif s == "3.5 数据可视化模块设计与实现" and "设计与实现" in s:
            # 第二个「3.5 数据可视化模块设计与实现」在 duplicate block — 已改为 3.9.2
            pass

    for para in doc.paragraphs:
        if para.text.strip() == "3.5 数据可视化模块设计与实现":
            # 仅替换 duplicate section：位于 3.9.1 之后的那一处；通过上下文简单处理——两处同名都改成 3.9.2 会冲突
            pass

    # 将第二次出现的「3.5 数据可视化模块设计与实现」（紧跟在 3.9.1 分析模块之后）改为 3.9.2
    seen_visual = 0
    for para in doc.paragraphs:
        if para.text.strip() == "3.5 数据可视化模块设计与实现":
            seen_visual += 1
            if seen_visual >= 2:
                para.clear()
                para.add_run("3.9.2 可视化与异步接口在代码中的落点")
                para.style = "Heading 3"
                break

    replace_first(
        lambda t: "4.1." in t and "测试环境" in t,
        "4.1 测试环境与数据概况",
        doc,
    )

    # ---- 在 4.6 之前插入 4.6~4.8 新内容，原 4.6 改为 4.9 ----
    target = None
    for para in doc.paragraphs:
        if para.text.strip().startswith("4.6 分析结果讨论"):
            target = para
            break
    if target is not None:
        para = target
        para.clear()
        para.add_run("4.9 综合讨论与局限")
        para.style = "Heading 2"

        # 在修改后的段落之前插入一系列段落（从后往前插，保持顺序）
        insert_paragraph_before(
            para,
            "复杂查询可在 SQL 层完成多表联结与时间窗口聚合；情感打分依赖应用层循环调用 SnowNLP，缺少库内「情感标签列」时，"
            "大规模统计宜采用批量导出 + 离线模型或定时任务物化结果表，以弥补当前「交互分析为主、重型 OLAP 为辅」的短板。",
        )
        insert_paragraph_before(para, "短板方面：每次请求新建 PyMySQL 连接，高并发下连接开销明显；"
                                     "复杂条件筛选（如跨分区、多关键词布尔检索）主要依赖 application 内存中的 pandas/jieba，而非数据库全文索引；"
                                     "评论情感未落库，重复分析会重复计算。")
        insert_paragraph_before(para, "示例：按 crawl_time 窗口 JOIN ranking_records 与 videos 取榜单快照对应的播放与分区分布；"
                                     "按 tname 分组聚合 AVG(view_count)；对 comments 按 bvid 聚合 COUNT 并与 videos 联结得到互动强度。"
                                     "此类查询可作为「热点演变」与「分区对比」报表的 SQL 基础。")
        insert_paragraph_before(para, "数据库采用 bilibili_hotspot（utf8mb4），核心表 videos、ranking_records、comments 由 db_operations 读写；"
                                     "查询入口包括 get_top_videos、query_videos、query_all_comments、get_video_trend 等，对应排行榜、关键词语料与趋势折线。")
        insert_paragraph_before(para, "4.8 数据库系统架构、短板与复杂分析查询", "Heading 2")

        insert_paragraph_before(
            para,
            "实验观察到：TF-IDF 更抑制通用高频词，有利于突出区分性话题词；词频更直观但易受「视频、哈哈」类词支配。"
            "同一时间段内对标题与评论分别提取时，标题侧重议题命名，评论侧重互动用语，二者 Top-K 重叠率可作为「话题一致性」旁证。",
        )
        insert_paragraph_before(para, "在固定时间窗与数据源（标题/评论）下，同时运行 analyze_keywords 与 analyze_keywords_tfidf，对比 Top-30 重叠率与互补词项；"
                                     "该对照用于验证第 3 章所述两条技术路线在实际语料上的差异，而非重新训练分词模型。")
        insert_paragraph_before(para, "词频法：分词后计数，反映语言表层共现；TF-IDF： Term Frequency × Inverse Document Frequency，"
                                     "降低全体文档常见词权重、抬高区分度词权重（实现为 jieba.analyse.extract_tags）。")
        insert_paragraph_before(para, "4.7 关键词提取方法对比与进一步验证", "Heading 2")

        insert_paragraph_before(
            para,
            "该评估用于论证 SnowNLP 在本数据集上的可用边界：即便总体分布合理，仍应对误判类型（讽刺、反语、混合情感句）保持审慎解释。",
        )
        insert_paragraph_before(para, "在标注集上可计算总体准确率 Accuracy = (TP+TN+…)/N；对「正面」一类可报告 Precision = TP/(TP+FP)。"
                                     "多分类可采用宏平均以观察各类是否均衡被识别。")
        insert_paragraph_before(para, "结果比对：以人工标签为 gold standard，与 SnowNLP 阈值分类结果逐条对照，统计混淆矩阵；"
                                     "可报告「中性类最易与相邻类混淆」等现象，与第 4.5 节的占比描述相互印证。")
        insert_paragraph_before(para, "抽样与标注：从 comments 表随机分层抽取不少于 200 条（可按正负中性初步占比分层），"
                                     "由标注员按同一细则打三类标签； disagreements 通过第三人仲裁，保证可复现。")
        insert_paragraph_before(para, "情感模型在本系统中指 SnowNLP 输出的连续得分经阈值离散化后的三分类结果。"
                                     "记准类样本上的准确率（Accuracy）为正确条数占比；对单一类别（如正面）可计算精准率（Precision）= 该类别预测且确实为该类的比例 / 该类别预测总数。")
        insert_paragraph_before(para, "4.6 情感模型有效性评估与标注实验", "Heading 2")

    out_path = path
    doc.save(str(out_path))
    print("OK:", out_path)


if __name__ == "__main__":
    main()
