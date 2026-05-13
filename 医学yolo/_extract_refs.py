# -*- coding: utf-8 -*-
"""Locate 参考文献 section and print entries (UTF-8 stdout)."""
import sys

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

p = r"医学骨架30012.docx"
d = Document(p)

# Search backwards for section heading
start = None
for i in range(len(d.paragraphs) - 1, -1, -1):
    t = (d.paragraphs[i].text or "").strip()
    if t == "参考文献" or t.startswith("参考文献"):
        start = i
        print("START index:", i, repr(t))
        break

if start is None:
    for i, para in enumerate(d.paragraphs):
        t = (para.text or "").strip()
        if "参考文献" in t and len(t) < 40:
            print("Fuzzy START", i, repr(t))
            start = i
            break

if start is not None:
    for j in range(start, len(d.paragraphs)):
        t = d.paragraphs[j].text.strip()
        if t:
            safe = t.replace("\u2212", "-")
            print(j, safe[:500])
else:
    print("No start found")
