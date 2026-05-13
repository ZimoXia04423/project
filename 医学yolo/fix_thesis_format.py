# -*- coding: utf-8 -*-
"""Apply 6 formatting fixes to thesis docx; write *_格式修订.docx."""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

SRC = Path(r"d:\project\医学python\基于深度学习的骨折检测系统设计与实现(2).docx")
DST = Path(r"d:\project\医学python\基于深度学习的骨折检测系统设计与实现(2)_格式修订.docx")


def add_paragraph_bottom_border(paragraph, sz: str = "12", color: str = "808080") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    old = p_pr.find(qn("w:pBdr"))
    if old is not None:
        p_pr.remove(old)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_style = p_pr.find(qn("w:pStyle"))
    jc = p_pr.find(qn("w:jc"))
    if p_style is not None:
        p_style.addnext(p_bdr)
    elif jc is not None:
        jc.addprevious(p_bdr)
    else:
        p_pr.insert(0, p_bdr)


def set_heading_spacing_lines(paragraph, before_lines: int | None, after_lines: int | None) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            p_style.addnext(spacing)
        else:
            p_pr.insert(0, spacing)
    if before_lines is not None:
        spacing.set(qn("w:beforeLines"), str(before_lines))
    if after_lines is not None:
        spacing.set(qn("w:afterLines"), str(after_lines))
    spacing.set(qn("w:line"), "440")
    spacing.set(qn("w:lineRule"), "auto")


def insert_caption_after_image(doc: Document, next_paragraph_index: int, caption_text: str) -> None:
    """Insert a centered caption immediately before paragraph at next_paragraph_index."""
    target = doc.paragraphs[next_paragraph_index]
    cap = target.insert_paragraph_before(caption_text)
    cap.style = "Thesis Body"
    cap.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap.paragraph_format.space_before = None
    cap.paragraph_format.space_after = None
    cap.paragraph_format.line_spacing = 1.0
    for r in cap.runs:
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = None


def apply_header_line(doc: Document) -> None:
    for sec in doc.sections:
        hdr = sec.header
        for p in hdr.paragraphs:
            t = (p.text or "").strip()
            if "毕业论文" in t or "设计" in t:
                add_paragraph_bottom_border(p)
                break


def apply_formula_number(doc: Document, para_index: int, number: str) -> None:
    p = doc.paragraphs[para_index]
    formula_text = p.text.strip()
    if not formula_text:
        return
    for r in list(p.runs):
        p._element.remove(r._element)
    p_pr = p._element.get_or_add_pPr()
    old_tabs = p_pr.find(qn("w:tabs"))
    if old_tabs is not None:
        p_pr.remove(old_tabs)
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center")
    tab_c.set(qn("w:pos"), str(int(Cm(8).twips)))
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right")
    tab_r.set(qn("w:pos"), str(int(Cm(15.5).twips)))
    tabs.extend([tab_c, tab_r])
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is not None:
        p_style.addnext(tabs)
    else:
        p_pr.insert(0, tabs)
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        p_pr.append(jc)
    else:
        jc.set(qn("w:val"), "left")

    p.add_run().add_tab()
    p.add_run(formula_text)
    p.add_run().add_tab()
    p.add_run(number)


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    if not SRC.is_file():
        print("Missing source:", SRC)
        return 1
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    apply_header_line(doc)

    # 图二：连续标题 1.2 / 1.2.1 减少段前段后叠加
    set_heading_spacing_lines(doc.paragraphs[117], before_lines=None, after_lines=0)
    set_heading_spacing_lines(doc.paragraphs[118], before_lines=0, after_lines=None)

    # 图四：公式序号（2.4 节首条独立公式）
    apply_formula_number(doc, 197, "(2-1)")

    # 图三、五、六：自后向前插入，避免段落索引错位（182=Neck，260=3.4 架构，292=4.1.2 影像）
    insert_caption_after_image(doc, 292, "图4-1  腕部 X 光影像示例")
    insert_caption_after_image(doc, 260, "图3-1  系统总体分层架构图")
    insert_caption_after_image(doc, 182, "图2-1  YOLOv10 Neck 多尺度特征融合结构示意图")

    doc.save(str(DST))
    print("Written:", DST)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
