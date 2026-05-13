# -*- coding: utf-8 -*-
"""
在 医学骨架300.docx 中：去掉标题里的占位批注，并在对应小节下插入架构图与预处理/增强效果图。

插图放入与模板一致的 Word 原生表格（继承文档中已有表格的 w:tblStyle，通常为 Normal Table / 样式 ID 20），
便于与学校模板表格线宽、边距一致；思路对齐 MiniMax minimax-docx「用 OpenXML 语义排版」而非纯浮动图片。
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def strip_placeholders(text: str) -> str:
    """移除红色占位批注（全角/半角括号）。"""
    patterns = [
        r"（系统架构图）",
        r"\(系统架构图\)",
        r"（制作的效果图）",
        r"\(制作的效果图\)",
    ]
    for pat in patterns:
        text = re.sub(pat, "", text)
    return re.sub(r"\s{2,}", " ", text).strip()


def get_thesis_table_style_name(doc: Document) -> str | None:
    """读取模板里第一张表格使用的样式名（如 Normal Table），用于新插入表格继承 tblStyle。"""
    for table in doc.tables:
        try:
            return table.style.name
        except (AttributeError, KeyError, ValueError):
            continue
    return None


def insert_table_after(paragraph: Paragraph, doc: Document, rows: int = 1, cols: int = 1) -> Table:
    """在段落之后插入表格（从文档末尾剪切到目标位置）。"""
    t = doc.add_table(rows=rows, cols=cols)
    tbl = t._tbl
    parent_container = t._parent
    tbl.getparent().remove(tbl)
    paragraph._element.addnext(tbl)
    return Table(tbl, parent_container)


def insert_paragraph_after_table(table: Table, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_figure_in_thesis_table(
    paragraph: Paragraph,
    doc: Document,
    image_path: Path,
    width_cm: float,
    style_name: str | None,
) -> Table:
    """单格表居中插图；返回表格对象，便于其后接图注段落。"""
    table = insert_table_after(paragraph, doc, rows=1, cols=1)
    if style_name:
        try:
            table.style = style_name
        except (KeyError, ValueError):
            pass
    cell = table.rows[0].cells[0]
    p_cell = cell.paragraphs[0]
    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cell.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return table


def main():
    root_np = Path(__file__).resolve().parent.parent / "newprogram"
    doc_src = root_np / "医学骨架300.docx"
    if not doc_src.is_file():
        raise SystemExit(f"找不到: {doc_src}")

    fig_fw = root_np / "figures_framework"
    arch_img = fig_fw / "framework_system_architecture.png"
    mod_img = fig_fw / "framework_module_structure.png"
    if not arch_img.is_file() or not mod_img.is_file():
        raise SystemExit(
            "请先运行 newprogram/generate_framework_figures.py 生成 figures_framework 下的 PNG。"
        )

    viz_dir = root_np / "visualization_results"
    pre_jpgs = sorted(viz_dir.glob("preprocessing_effect_*.jpg"))
    aug_jpgs = sorted(viz_dir.glob("augmentation_effect_*.jpg"))
    if not pre_jpgs or not aug_jpgs:
        raise SystemExit(f"请先运行 newprogram/visualize_effects.py，缺少: {viz_dir}")

    pre_img = pre_jpgs[-1]
    aug_img = aug_jpgs[-1]

    dst = root_np / "医学骨架300_插图版.docx"
    try:
        shutil.copy2(doc_src, dst)
    except PermissionError:
        dst = root_np / "医学骨架300_插图表格版.docx"
        shutil.copy2(doc_src, dst)
        print("提示: 原插图版文件被占用，已写入:", dst.name)
    doc = Document(str(dst))
    tbl_style = get_thesis_table_style_name(doc)

    # 标注 -> (图片路径, 图注)
    inserted_arch = False
    inserted_mod = False
    inserted_pre = False
    inserted_aug = False

    paras = list(doc.paragraphs)
    for para in paras:
        raw = para.text
        if not raw.strip():
            continue

        # 3.4 架构：标题去批注 + 插图
        if ("3.4" in raw and "系统总体架构" in raw) or ("系统架构图" in raw and "3.4" in raw):
            cleaned = strip_placeholders(raw)
            if cleaned != raw:
                para.clear()
                para.add_run(cleaned)
            if not inserted_arch and arch_img.is_file():
                fig_tbl = insert_figure_in_thesis_table(para, doc, arch_img, 14.0, tbl_style)
                insert_paragraph_after_table(
                    fig_tbl,
                    "图 3-1  系统总体架构示意图（四层结构及典型业务流程，见生成稿）。",
                )
                inserted_arch = True
            continue

        if "3.5" in raw and "模块" in raw and "设计" in raw:
            cleaned = strip_placeholders(raw)
            if cleaned != raw:
                para.clear()
                para.add_run(cleaned)
            if not inserted_mod and mod_img.is_file():
                fig_tbl = insert_figure_in_thesis_table(para, doc, mod_img, 14.0, tbl_style)
                insert_paragraph_after_table(
                    fig_tbl,
                    "图 3-2  系统模块结构示意图（datasets / preprocess / models / train / infer / gui）。",
                )
                inserted_mod = True
            continue

        if "4.1.2" in raw and "预处理" in raw:
            cleaned = strip_placeholders(raw)
            if cleaned != raw:
                para.clear()
                para.add_run(cleaned)
            if not inserted_pre:
                fig_tbl = insert_figure_in_thesis_table(para, doc, pre_img, 14.0, tbl_style)
                insert_paragraph_after_table(
                    fig_tbl,
                    "图 4-1  图像预处理流程效果示意（灰度归一化、CLAHE、去噪与尺寸标准化等步骤）。",
                )
                inserted_pre = True
            continue

        if "4.1.3" in raw and "增强" in raw:
            cleaned = strip_placeholders(raw)
            if cleaned != raw:
                para.clear()
                para.add_run(cleaned)
            if not inserted_aug:
                fig_tbl = insert_figure_in_thesis_table(para, doc, aug_img, 14.0, tbl_style)
                insert_paragraph_after_table(
                    fig_tbl,
                    "图 4-2  数据增强效果示意（随机翻转、旋转、亮度/对比度扰动与缩放平移等）。",
                )
                inserted_aug = True
            continue

    doc.save(str(dst))
    print("已生成:", dst)
    print("  继承表格样式:", tbl_style)
    print("  架构图:", arch_img)
    print("  模块图:", mod_img)
    print("  预处理:", pre_img.name)
    print("  增强: ", aug_img.name)
    print("插入标记: arch=", inserted_arch, " mod=", inserted_mod, " pre=", inserted_pre, " aug=", inserted_aug)


if __name__ == "__main__":
    main()
