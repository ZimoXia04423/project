# -*- coding: utf-8 -*-
"""
在「医学骨架30012.docx」参考文献列表末尾追加可核查的中文文献条目。
跳过条件：已存在「10.15953/j.1004-4140.2022.31.05.08」（首条中文 DOI）则不再写入。
"""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


MARKER_DOI = "10.15953/j.1004-4140.2022.31.05.08"

# GB/T 7714 著录格式；卷期页与 DOI 可独立核查（期刊官网 / Europe PMC / 知网）
REF_ZH = [
    "[30] 熊山, 陈博, 毛杰, 刘四斌, 黄原义, 程建敏. 基于深度学习的计算机辅助诊断系统在肋骨骨折诊断中的应用[J]. CT理论与应用研究, 2022, 31(5): 617-622. DOI: 10.15953/j.1004-4140.2022.31.05.08.",
    "[31] 田冲, 陈新, 朱华, 秦松, 石亮, 芮永. 机器学习在创伤骨科中的应用与展望[J]. 中国修复重建外科杂志, 2023, 37(12): 1562-1568. DOI: 10.7507/1002-1892.202308064. (PMID: 38130202, PMC: PMC10739668)",
    "[32] 谭辉, 田占雨, 潘宁, 等. 基于深度学习的计算机辅助诊断系统在提高急性肋骨骨折诊断效能上的价值[J]. 临床放射学杂志, 2020, 39(12): 2493-2497. DOI: 10.13437/j.cnki.jcr.2020.12.030.",
]


def patch(path: Path) -> bool:
    doc = Document(str(path))
    full = "\n".join(p.text for p in doc.paragraphs)
    if MARKER_DOI in full:
        print("skip: Chinese refs already present")
        return False

    anchor = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("[29]"):
            anchor = p
            break
    if anchor is None:
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith("[28]"):
                anchor = p
        if anchor is None:
            print("error: no [29] or [28] ref paragraph found")
            return False

    cur = anchor
    for line in REF_ZH:
        cur = insert_paragraph_after(cur, line)

    out = path
    try:
        doc.save(str(out))
    except PermissionError:
        out = path.with_name(path.stem + "_已加中文文献.docx")
        doc.save(str(out))
        print("saved (file was locked):", out)
        return True

    print("saved:", out)
    return True


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    patch(root / "医学骨架30012.docx")
