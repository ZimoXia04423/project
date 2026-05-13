# -*- coding: utf-8 -*-
"""
将第 6 章「总结与展望」替换为更详实、并与当前桌面检测系统实现相呼应的正文。
定位：首段含「总结与展望」且以 6 起头的章节标题 -> 下一节「参考文献」之前。
"""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


MARKER = "FractureGrader后置融合分级器"

SECTION_6 = [
    "6 总结与展望",
    "6.1 全文总结",
    "本文围绕「基于深度学习的骨折检测系统设计与实现」，面向腕部 X 光影像中骨折线细小、对比度低、正负样本严重失衡以及算法成果难以形成可演示软件等问题，完成了从数据与模型到桌面端系统的贯通式工作。整体上，研究路径与工程路径相互支撑：前者给出可写、可评、可比的实验设计，后者将推理结果以可视化方式固化到交互流程中，便于答辩展示与后续迭代。",
    "在模型与数据层面，本文以 YOLO 系列实时检测为技术主线，结合公开腕部创伤数据集组织思路，完成面向 YOLO 训练的样本整理、划分与增强，并保留 YOLOv8、YOLOv10 等训练日志与权重产物，使「训练—验证—测试」链条可追溯。论文层面所述的坐标注意力与 Focal Loss 等改进，为提升细粒度特征与缓解难易样本失衡提供了理论依据；工程侧则以 Ultralytics 框架加载最优权重完成推理，为系统稳定运行提供现成接口。",
    "在系统实现层面，本文完成了基于 PyQt5 的桌面端「骨折分级智能检测系统」原型：程序入口完成 Fusion 风格与字体初始化后，先通过本地 JSON 账户文件完成登录/注册，再进入主界面。主界面采用无边框医疗科技风布局，左侧展示原始影像，中部展示带框标注结果，右侧集中呈现参数与结论。推理阶段在独立线程中调用 YOLO，避免阻塞界面；支持单模型检测与 YOLOv8/YOLOv10 双模型对比，并可调节置信度阈值与 IoU，以满足不同演示与灵敏度需求。",
    "在业务逻辑层面，系统将模型输出的多类别病灶框与面向临床表述的「骨折分级」解耦：检测头输出骨折线、骨膜反应、旋前征、软组织肿胀等细粒度类别（与数据 yaml 中定义一致），再由 FractureGrader 后置融合分级器依据「是否检出骨折及伴随征象」的规则链，自动给出重度、中度、轻度或未见骨折的结论，并在界面卡片与状态栏同步展示。该设计使深度模型专注定位与分类，使分级策略可解释、可单测、可随指南调整，有利于后续与临床专家共识对齐。",
    "在评价与文档层面，本文整理了功能测试与性能测试思路，给出模型对比与消融实验框架，并从实时性、可维护性、可扩展性等角度对系统需求作出回应。需要客观指出的是：当前分级规则尚未融合置信度加权与多框投票；系统定位为科研与教学演示原型，不等同于获批医疗器械软件；跨机构、跨设备数据分布差异亦需在更大样本上进一步验证。",
    "总体来看，本文在「算法方案 + 工程原型」两个维度形成了较完整成果：既为腕部骨折影像智能辅助筛查提供了可继续打磨的技术路线，也以可运行软件形式把模型能力落到具体交互闭环中，为后续外部验证、轻量化部署与规范化文档编写打下了基础。",
    "6.2 研究展望",
    "（1）数据与标注体系：在合规与脱敏前提下，持续扩充多中心、多设备、多年龄段腕部 X 光样本；探索半监督、主动学习与难例回灌机制，降低标注成本；完善患者级去重与外部测试集，形成可发表的严格验证流程。",
    "（2）模型与训练：在保持实时性的前提下，将论文中的结构改进与损失设计与训练代码深度对齐并做系统消融；研究将置信度、框重叠度纳入分级或二次筛选；引入校准（calibration）与错误模式分析，减少高置信误报对临床信任的伤害。",
    "（3）系统与工程：在现有 PyQt 桌面架构上扩展批量推理、历史记录与结果检索；支持 DICOM 导入与匿名化元数据清洗；提供安装包、自动更新与日志脱敏；将推理链路封装为本地 HTTP/gRPC 服务，便于与院内测试环境对接。",
    "（4）部署与性能：开展 ONNX/TensorRT 转换与 INT8 量化试验，评估在消费级 GPU 乃至 CPU 上的延迟与显存占用；针对高分辨率原图探索滑窗或金字塔推理，在精度与速度之间取得可配置折中。",
    "（5）临床与产品化路径：若走向真实辅助诊断，需另行规划医疗器械软件生存周期、临床试验或真实世界研究设计，并在界面显著位置提示「仅供研究/教学，不能替代医师诊断」；探索与 PACS、结构化报告或危急值提醒流程的衔接，但仍以安全与责任边界为首要约束。",
    "（6）可解释性与多模态：引入 Grad-CAM、注意力热图或基于 bbox 的区域对比，向使用者展示模型关注区域；在隐私合规基础上，探索融合侧位片、对侧对比或电子病历文本的多模态提示，以缓解单张正位片结构重叠导致的隐匿骨折漏检风险。",
]


def patch_document(path: Path) -> bool:
    doc = Document(str(path))
    if any(MARKER in p.text for p in doc.paragraphs):
        print("skip (chapter 6 already patched):", path.name)
        return False

    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if start is None and "总结与展望" in t and (t.startswith("6") or t.startswith("第6")):
            start = i
            continue
        if start is not None and end is None and t.startswith("参考文献"):
            end = i
            break
    if start is None or end is None:
        print("skip (chapter 6 boundaries not found):", path.name)
        return False

    while True:
        if start + 1 >= len(doc.paragraphs):
            break
        p_next = doc.paragraphs[start + 1]
        if p_next.text.strip().startswith("参考文献"):
            break
        delete_paragraph(p_next)

    first = doc.paragraphs[start]
    first.clear()
    cur = first
    for idx, block in enumerate(SECTION_6):
        if idx == 0:
            cur.add_run(block)
        else:
            cur = insert_paragraph_after(cur, block)

    for p in doc.paragraphs:
        rt = p.text.strip()
        if rt.startswith("参考文献"):
            if "要求" in rt or "篇" in rt:
                p.clear()
                p.add_run("参考文献")
            break

    try:
        doc.save(str(path))
    except PermissionError:
        alt = path.with_name(path.stem + "_第6章已修订.docx")
        doc.save(str(alt))
        print("saved (locked):", alt)
        return True
    print("patched chapter 6:", path)
    return True


def main():
    root = Path(__file__).resolve().parent.parent
    for name in (
        "基于深度学习的骨折检测系统设计与实现.docx",
        "基于深度学习的骨折检测系统设计与实现_完善版.docx",
    ):
        p = root / name
        if p.is_file():
            patch_document(p)


if __name__ == "__main__":
    main()
