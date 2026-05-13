# -*- coding: utf-8 -*-
"""
按导师意见：纵向对比仅保留「一种」YOLOv8，去掉并列的 YOLOv5 对照及表中对应行，
避免实验过多、难以完成。
"""
from pathlib import Path

from docx import Document


def patch_paragraphs(doc: Document) -> int:
    replacements = [
        (
            "为验证改进YOLOv10的有效性，应与如下模型进行比较：1. YOLOv8",
            "为验证改进 YOLOv10 的有效性，纵向对比中在同一代码框架与相同数据划分下，仅选取 "
            "YOLOv8 作为同系列基线（固定为一种网络规格与训练配置，不并列 YOLOv8-n/s/m "
            "等多种规格），以控制实验工作量并与改进模型公平对比。",
        ),
        (
            "对比重点包括Precision、Recall、mAP和推理时间。",
            "对比重点包括 Precision、Recall、mAP 与推理时间等指标。",
        ),
        (
            "在相同训练集、验证集和测试集条件下，本文分别训练了Faster R-CNN、SSD、YOLOv5、YOLOv8、原生YOLOv10和改进YOLOv10模型，并在测试集上进行统一评估，结果如表5-3所示。",
            "在相同训练集、验证集与测试集划分条件下，本文分别训练并评估了 Faster R-CNN、SSD、"
            "YOLOv8（仅采用一种固定规格）、原生 YOLOv10 与改进 YOLOv10，并在测试集上进行统一评估，结果如表 5-3 所示。",
        ),
        (
            "与YOLOv8和YOLOv5相比，改进YOLOv10在Precision、Recall和mAP指标上仍具有一定优势，表明YOLOv10端到端检测框架在医学影像骨折检测任务中具有较好的结构适应性。与Faster R-CNN相比，改进模型虽然在精度上仅有小幅优势，但推理速度明显更快，更适合实际辅助诊断场景下对实时性的要求。SSD在速度方面较快，但精度指标落后较多，不适合作为本文系统的最终部署模型。",
            "与 YOLOv8 相比，改进 YOLOv10 在 Precision、Recall 与 mAP 指标上仍具有一定优势，"
            "表明 YOLOv10 端到端检测框架在医学影像骨折检测任务中具有较好的结构适应性。"
            "与 Faster R-CNN 相比，改进模型虽然在精度上仅有小幅优势，但推理速度明显更快，"
            "更适合实际辅助诊断场景下对实时性的要求。SSD 在速度方面较快，但精度指标落后较多，"
            "不适合作为本文系统的最终部署模型。",
        ),
    ]
    n = 0
    for p in doc.paragraphs:
        t = p.text
        for old, new in replacements:
            if old in t:
                p.text = t.replace(old, new)
                n += 1
                break
    return n


def remove_yolov5_row(doc: Document) -> bool:
    """表「模型 / Precision / …」中删除 YOLOv5 所在行。"""
    target_tbl = None
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        if (tbl.rows[0].cells[0].text.strip() == "模型"
                and "Precision" in tbl.rows[0].cells[1].text):
            target_tbl = tbl
            break
    if target_tbl is None:
        print("warn: 未找到表 5-3 对应表格，跳过删行")
        return False
    row_idx = None
    for i, row in enumerate(target_tbl.rows):
        if row.cells[0].text.strip() == "YOLOv5":
            row_idx = i
            break
    if row_idx is None:
        print("warn: 未找到 YOLOv5 行，跳过删行")
        return False
    tr = target_tbl.rows[row_idx]._tr
    tr.getparent().remove(tr)
    print(f"已删除表格中第 {row_idx} 行（YOLOv5）")
    return True


def main():
    root = Path(__file__).resolve().parents[1]
    path = root / "医学骨架30012.docx"
    doc = Document(str(path))
    n = patch_paragraphs(doc)
    print(f"已替换段落 {n} 处（预期 4 处）")
    remove_yolov5_row(doc)
    try:
        doc.save(str(path))
    except PermissionError:
        alt = path.with_name(path.stem + "_4.7已改.docx")
        doc.save(str(alt))
        print("原文档占用，已保存：", alt)
        return
    print("已保存：", path)


if __name__ == "__main__":
    main()
