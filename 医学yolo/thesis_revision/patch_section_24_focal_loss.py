# -*- coding: utf-8 -*-
"""
将论文中「2.4 Focal Loss」小节替换为带公式、符号说明与医学检测语境的完整正文。
通过段落标题「2.4…Focal」与下一节「2.5」定位范围。
"""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


# 小节正文：分段便于 Word 排版（可在 Word 中再调首行缩进/段距）
SECTION_24_BODY = [
    "在目标检测任务中，锚框或网格单元在数量上往往远大于真实目标框，导致训练样本在类别上呈现极端不平衡：大量「背景」或「易分类负样本」在每次迭代中产生稳定、较小的损失，从而在梯度层面压制了少量「前景」或「难分类正样本」的贡献。若仍采用常规交叉熵作为分类损失，模型容易过早收敛到对背景预测「过于自信」的状态，进而削弱对细小、低对比度骨折线等难例的学习能力。",
    "记二分类下真实标签 y∈{0,1}，模型对正类的预测概率为 p∈(0,1)。为统一书写，定义 p_t：当 y=1 时 p_t=p；当 y=0 时 p_t=1-p。则交叉熵可写为 CE(p_t)=−ln(p_t)。当分类「很容易」时 p_t→1，CE→0，该样本对梯度的贡献迅速衰减；但在不平衡设定下，海量易分类样本的微小梯度叠加后，仍可能主导参数更新方向。",
    "为缓解上述问题，Lin 等人在 RetinaNet 中提出 Focal Loss（FL），在交叉熵外引入调制因子 (1−p_t)^γ，使易分类样本（p_t 接近 1）的损失被进一步压缩，从而使优化过程更关注 p_t 较小的难分类样本。其常见形式为：",
    "FL(p_t) = −α_t · (1−p_t)^γ · ln(p_t)。",
    "其中：p_t 为前述「真实类别上的预测概率」；α_t∈(0,1) 为类别平衡因子，用于缓解正负（或多类）样本数量差异，实践中可对正、负类取不同 α 或对稀有类赋予更大权重；γ≥0 为聚焦参数，γ=0 时退化为加权交叉熵；γ 越大，易分类样本被抑制的程度越强。目标检测与实例分割中常取 γ=2 作为经验起点，并在验证集上微调。",
    "从函数形态看，当 p_t→1 时，(1−p_t)^γ 以幂次速度趋于 0，易分类样本对总损失的贡献被显著削弱；当 p_t 较小（难例）时，(1−p_t)^γ 接近常数阶，难例仍保持与 CE 同量级的梯度强度，从而把有限的模型容量「推」向决策边界附近的困难区域。",
    "在多类别检测中，可在每个锚框/每个网格位置的分类 logits 上对各类别计算焦点调制后的损失再求和或取平均，并与边界框回归、目标性（objectness）等损失加权组合，形成端到端训练目标。需注意：Focal Loss 主要作用于分类分支，对框回归仍常采用 Smooth L1 / IoU 系列损失；与在线难例挖掘（OHEM）等方法相比，FL 以连续权重替代离散筛选，实现简单且与 mini-batch 训练兼容。",
    "在腕部 X 光骨折检测场景中，骨折线往往仅占整幅图像的极小面积，且与骨皮质、骨小梁纹理在灰度上高度相似，前景「正」与背景「负」在像素与候选框层面均严重失衡；同时，大量「看似可疑但实为正常解剖变异」的区域构成难负样本。将 Focal Loss 思想引入分类分支，有助于降低简单背景对梯度的支配作用，提升对真正骨折候选及难例的关注度，从而在一定程度上改善召回率与 mAP，并与坐标注意力等结构增强手段形成互补。",
    "需要说明的是，γ 与 α 的选取与数据分布、批大小及学习率等强相关，过大 γ 可能导致训练早期梯度过于稀疏；实际系统中应结合验证集曲线与消融实验确定超参数，并配合采样策略、数据增强与阈值后处理共同优化整体检测性能。",
]


def patch_document(path: Path) -> bool:
    doc = Document(str(path))
    if any("锚框或网格单元在数量上往往远大于真实目标框" in p.text for p in doc.paragraphs):
        print("skip (already patched):", path.name)
        return False
    paras = doc.paragraphs
    start = end = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if start is None and t.startswith("2.4") and "Focal" in t:
            start = i
            continue
        if start is not None and end is None and t.startswith("2.5"):
            end = i
            break
    if start is None or end is None:
        print("skip (section not found):", path.name)
        return False

    # 删除 2.4 标题之后到 2.5 之前的所有段落
    while True:
        if start + 1 >= len(doc.paragraphs):
            break
        p_next = doc.paragraphs[start + 1]
        if p_next.text.strip().startswith("2.5"):
            break
        delete_paragraph(p_next)

    title_para = doc.paragraphs[start]
    title_para.clear()
    title_para.add_run("2.4 Focal Loss")

    cur = title_para
    for block in SECTION_24_BODY:
        cur = insert_paragraph_after(cur, block)

    try:
        doc.save(str(path))
    except PermissionError:
        alt = path.with_name(path.stem + "_2.4已修订.docx")
        doc.save(str(alt))
        print("saved (original locked):", alt)
        return True
    print("patched:", path)
    return True


def main():
    root = Path(__file__).resolve().parent.parent
    candidates = [
        root / "基于深度学习的骨折检测系统设计与实现.docx",
        root / "基于深度学习的骨折检测系统设计与实现_完善版.docx",
    ]
    for p in candidates:
        if p.is_file():
            try:
                patch_document(p)
            except PermissionError as e:
                print("permission:", p.name, e)


if __name__ == "__main__":
    main()
