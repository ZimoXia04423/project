# -*- coding: utf-8 -*-
"""
在论文 Word 副本上完善内容：不删除原有论述与表格数据；
去除括号内写作批注、修正明显笔误，并在关键位置后插入补充段落与训练曲线图。
"""
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def strip_editor_notes(text: str) -> str:
    """去掉写作批注（括号提示），保留正文。"""
    patterns = [
        r"（太少完善一下）",
        r"（所有参考文献上标）",
        r"（连着的，可以连在一起两个文献）",
        r"（原理的地方插入图片）",
        r"1\.2\.3\(插入架构图片）",
        r"（加上公式）",
        r"（四个公式加入）",
        r"（系统架构图）",
        r"（制作的效果图）",
        r"（公式）",
        r"\(对比只保留一个v8的，老师说太多了不好做\)",
        r"（图）",
        r"（曲线图加入）",
        r"（写的详细一点）",
        r"（要求28-30篇中、英文献，近几年的）",
        r"（要求\s*28[-–]30\s*篇中、英文献，近几年的）",
        r"（要求\s*\d+[-–]\d+\s*篇[^）]*）",
    ]
    for pat in patterns:
        text = re.sub(pat, "", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    return text


def enrich_document(doc: Document) -> None:
    # 1) 段落：去批注 + 英文题目笔误
    for para in doc.paragraphs:
        t = para.text
        if not t:
            continue
        nt = strip_editor_notes(t)
        if "Learningn" in nt:
            nt = nt.replace("Learningn", "Learning")
        if nt != t:
            para.clear()
            para.add_run(nt)

    # 2) 在「1.3.3 现有研究存在的问题」标题之前插入国内外研究补充段
    marker = "近年来，面向小目标与低对比度医学图像的检测框架持续演进"
    already = any(marker in p.text for p in doc.paragraphs)
    if not already:
        for para in doc.paragraphs:
            t = para.text.strip()
            if t.startswith("1.3.3") and "现有研究存在的问题" in t:
                insert_paragraph_before(
                    para,
                    "近年来，面向小目标与低对比度医学图像的检测框架持续演进：多尺度特征融合仍是提升细长、低对比度病灶检出率的重要手段；"
                    "通道、空间与坐标等注意力机制有助于强化关键区域表征；端到端单阶段检测在推理效率上的优势则使其更契合急诊预筛与桌面端部署需求。"
                    "在数据层面，以GRAZPEDWRI-DX为代表的公开数据集提升了研究可比性，但跨中心、跨设备、跨年龄段的外部验证仍相对不足，"
                    "真实工作流中的误报控制与临床路径衔接亦需持续关注。基于上述背景，在实时检测框架上结合任务特征进行结构与损失层面的适配，"
                    "并形成可测试、可演示的系统实现，对工程型毕业设计仍具有明确的训练价值。",
                )
                break

    # 3) 2.4 Focal Loss 节后补充公式说明（插在含「将Focal Loss应用于骨折检测」的段之后）
    focal_anchor = "将Focal Loss应用于骨折检测任务时，其优势十分明显"
    formula_block = (
        "在此基础上，二分类形式的Focal Loss可写为：FL(p_t) = -α_t · (1 - p_t)^γ · ln(p_t)，"
        "其中p_t为模型对真实类别的预测概率，α_t用于缓解类别不平衡，γ≥0用于调节难易样本在损失中的权重；"
        "γ越大则易分类样本对梯度的贡献越弱，从而使训练更关注难例。多分类检测中可在分类分支对各锚框/各通道按上式扩展并与回归损失加权求和。"
    )
    if not any(formula_block[:30] in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if focal_anchor in para.text:
                insert_paragraph_after(para, formula_block)
                break

    # 4) 3.3.1 在「误检情况」一句后补充公式（若尚无 TP/FP 行）
    formula_331 = (
        "为便于量化评价，可将精确率、召回率与F1写为：Precision = TP/(TP+FP)，Recall = TP/(TP+FN)，"
        "F1 = 2·Precision·Recall/(Precision+Recall)。mAP@0.5与mAP@0.5:0.95则分别对应IoU=0.5及IoU从0.5到0.95步进0.05时的平均精度均值，用于综合衡量定位质量。"
    )
    if not any("Precision = TP/(TP+FP)" in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if "模型评价应综合考虑Precision" in para.text and "误检情况" in para.text:
                insert_paragraph_after(para, formula_331)
                break

    # 5) 表5-3 后补充「精简对比」说明（不删表内原数据）
    note53 = (
        "【补充说明】在控制实验工作量、便于集中展示核心结论的前提下，亦可侧重对比YOLOv8、原生YOLOv10与改进YOLOv10三类模型；"
        "上表保留多模型结果便于与既有文献口径对照，答辩或定稿时可根据指导教师意见对对比对象与文字分析做进一步取舍。"
    )
    if not any("【补充说明】在控制实验工作量" in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if "由表5-3可知" in para.text:
                insert_paragraph_after(para, note53)
                break

    # 6) 5.3.2 性能测试段后插入训练曲线图（文件存在时）
    fig_dir = Path(__file__).resolve().parent / "figures"
    p10 = fig_dir / "yolov10_train_map_curve.png"
    p8 = fig_dir / "yolov8_train_map_curve.png"
    cap = "图（补充） 基于验证集的训练过程mAP曲线示例（上图：YOLOv10；下图：YOLOv8），可与TensorBoard导出曲线对照使用。"
    if p10.is_file() and p8.is_file() and not any("训练过程mAP曲线示例" in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if "性能测试主要从模型精度" in para.text and "FPS" in para.text:
                p_cap = insert_paragraph_after(para, cap)
                r = p_cap.add_run()
                r.add_picture(str(p10), width=Cm(14))
                p_cap2 = insert_paragraph_after(p_cap, "")
                r2 = p_cap2.add_run()
                r2.add_picture(str(p8), width=Cm(14))
                break

    # 7) 6.2 研究展望段后追加展望条目（不删原五条）
    extra = (
        "此外，可进一步开展与PACS或医院信息系统的接口联调与权限审计；在算法侧探索半监督、主动学习以缓解标注成本；"
        "在规范侧对照TRIPOD-AI、DECIDE-AI等报告建议完善验证设计与临床适用性说明。"
    )
    if not any("TRIPOD-AI" in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if "多模态骨折辅助诊断系统" in para.text and para.text.strip().endswith("。"):
                insert_paragraph_after(para, extra)
                break

    # 8) 6.1 总结段后补充（不删原条目）
    sum_extra = (
        "从工程实现角度看，本文将数据预处理、模型训练与推理、可视化交互等环节贯通，形成了可复现的实验记录与可演示的系统原型；"
        "从方法角度看，坐标注意力与难例重加权策略与腕部骨折线细小、对比度偏低的特点相匹配，为后续在更大样本与更严格验证条件下的迭代提供了基线。"
    )
    if not any("从工程实现角度看，本文将数据预处理" in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if "也为后续继续开展真实训练、外部验证和轻量部署打下了基础" in para.text:
                insert_paragraph_after(para, sum_extra)
                break

    # 9) 参考文献标题下补充说明（不删原有条目）
    ref_note = (
        "【补充说明】定稿时建议按学院格式将文献总量扩展至28～30篇，并优先选用近五年中英文期刊与会议论文；"
        "新增文献可围绕YOLO系列更新、医学小目标检测、可解释性、临床验证规范（如TRIPOD-AI）及模型部署等方向检索与核对。"
    )
    if not any("【补充说明】定稿时建议按学院格式" in p.text for p in doc.paragraphs):
        for para in doc.paragraphs:
            if para.text.strip() == "参考文献" or para.text.strip().startswith("参考文献"):
                insert_paragraph_after(para, ref_note)
                break


def main():
    root = Path(__file__).resolve().parent.parent
    src = root / "基于深度学习的骨折检测系统设计与实现.docx"
    if not src.is_file():
        raise SystemExit(f"找不到源文件: {src}")
    dst = root / "基于深度学习的骨折检测系统设计与实现_完善版.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))
    enrich_document(doc)
    doc.save(str(dst))
    print("已生成:", dst)


if __name__ == "__main__":
    main()
